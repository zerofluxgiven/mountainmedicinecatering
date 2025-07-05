import io
import os
import tempfile
import mimetypes
import csv
import re
import openai
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

# Check if Tesseract is available
try:
    tesseract_version = pytesseract.get_tesseract_version()
    print(f"Tesseract version: {tesseract_version}")
except Exception as e:
    print(f"WARNING: Tesseract not found or not configured: {e}")
from docx import Document
import requests
from bs4 import BeautifulSoup
import json
from datetime import datetime
import streamlit as st
from firebase_init import db, firestore
from utils import normalize_keys, normalize_recipe_quantities
from recipes import (
    save_recipe_to_firestore,
    save_event_to_firestore,
    save_menu_to_firestore,
    save_ingredient_to_firestore,
)

# Initialize OpenAI client
try:
    # Check for environment variable first (which might be overriding)
    env_key = os.getenv("OPENAI_API_KEY")
    if env_key:
        print(f"WARNING: OPENAI_API_KEY environment variable is set and starts with: {env_key[:10]}...")
    
    # Use secrets
    api_key = st.secrets["openai"]["api_key"]
    
    # The OpenAI library might be picking up an env var instead of our provided key
    # Force it to use our key by clearing any env var
    if "OPENAI_API_KEY" in os.environ:
        del os.environ["OPENAI_API_KEY"]
    
    client = openai.OpenAI(api_key=api_key)
except Exception as e:
    print(f"Failed to initialize OpenAI client: {e}")
    st.error(f"Failed to initialize OpenAI: {str(e)}")
    client = None

# --------------------------------------------
# 🧹 Central Cleaning & Validation Utils
# --------------------------------------------

def clean_raw_text(text: str) -> str:
    text = text.replace("\r", "").strip()
    lines = text.split("\n")
    cleaned = [line for line in lines if len(line.strip()) > 2 and not line.lower().startswith("http")]
    return "\n".join(cleaned)

def is_meaningful_recipe(recipe: dict) -> bool:
    if not isinstance(recipe, dict):
        return False
    if not (recipe.get("name") or recipe.get("title")):
        return False
    # Check if ingredients and instructions are not empty
    ingredients = recipe.get("ingredients", [])
    instructions = recipe.get("instructions", [])
    # Ensure we have actual content, not just empty lists
    has_ingredients = ingredients and len(ingredients) > 0
    has_instructions = instructions and len(instructions) > 0
    return has_ingredients and has_instructions

# --------------------------------------------
# 🧠 Main Entry Point (Patched)
# --------------------------------------------

def parse_file(uploaded_file, target_type="all", user_id=None, file_id=None):
    st.info("📄 Processing file...")
    print(f"📄 STARTING parse_file() - File: {getattr(uploaded_file, 'name', 'Unknown')}, Type: {getattr(uploaded_file, 'type', 'Unknown')}")

    raw_text = extract_text(uploaded_file)
    st.session_state["extracted_text"] = raw_text
    image_url = extract_image_from_file(uploaded_file)

    if raw_text and raw_text.strip():
        st.success(f"✅ Extracted {len(raw_text)} characters of text.")
        print(f"📄 Extracted text preview: {raw_text[:300]}...")
    else:
        st.error("❌ No text could be extracted from this file.")
        st.info("💡 Tip: Make sure the file contains readable text. For images, ensure they have clear text.")
        return {}

    parsed = {}
    target_types = [target_type] if target_type != "all" else [
        "recipes", "menus", "tags", "ingredients", "allergens"
    ]

    cleaned_text = clean_raw_text(raw_text)

    for t in target_types:
        parsed[t] = query_ai_parser(cleaned_text, t)
        if t == "recipes":
            recipe_data = parsed[t]
            if isinstance(recipe_data, dict) and "recipes" in recipe_data:
                recipe_data = recipe_data["recipes"]
            if isinstance(recipe_data, list):
                recipe_data = recipe_data[0] if recipe_data else {}
            recipe_data = normalize_keys(recipe_data)
            normalize_recipe_quantities(recipe_data)
            if image_url:
                recipe_data.setdefault("image_url", image_url)
            parsed[t] = recipe_data

    if file_id:
        parsed_record = {
            "parsed": parsed,
            "version": 1,
            "status": "unusable" if not is_meaningful_recipe(parsed.get("recipes", {})) else "pending_review",
            "raw_text": raw_text[:5000],
            "last_updated": datetime.utcnow(),
            "user_id": user_id
        }
        try:
            db.collection("files").document(file_id).update({"parsed_data": parsed_record})
            st.warning("✅ Saved parsed_data to Firestore")
        except Exception as e:
            print(f"Error saving parsed data to Firestore: {e}")

    return parsed

# --------------------------------------------
# 📄 Text Extraction
# --------------------------------------------

def extract_text(uploaded_file):
    mime_type, _ = mimetypes.guess_type(uploaded_file.name)
    
    # Debug info
    print(f"Extracting text from: {uploaded_file.name} (type: {uploaded_file.type}, mime: {mime_type})")
    
    try:
        uploaded_file.seek(0)
        
        # Check both file.type and mime_type for images
        if uploaded_file.type.startswith("image") or (mime_type and mime_type.startswith("image")):
            st.info(f"Processing image: {uploaded_file.name}")
            text = extract_text_from_image(uploaded_file)
            if text:
                st.success(f"Extracted {len(text)} characters from image")
            return text
        elif uploaded_file.type == "application/pdf":
            return extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type.startswith("text"):
            return uploaded_file.read().decode("utf-8")
        elif uploaded_file.type in ("text/csv", "application/vnd.ms-excel", "application/csv"):
            return extract_text_from_csv(uploaded_file)
        elif uploaded_file.type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return extract_text_from_docx(uploaded_file)
        else:
            # Fallback - check if it might be an image based on extension
            if uploaded_file.name.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
                st.info("Detected image by extension, attempting OCR...")
                return extract_text_from_image(uploaded_file)
            return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Text extraction failed: {str(e)}")
        print(f"Text extraction error: {e}")
        return ""

def extract_text_from_pdf(uploaded_file):
    text = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name
        doc = fitz.open(tmp_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        os.remove(tmp_path)
    except Exception as e:
        print(f"PDF parse error: {e}")
    return text

def extract_text_with_vision(uploaded_file):
    """Use OpenAI Vision API to extract text from image"""
    try:
        import base64
        uploaded_file.seek(0)
        
        # Encode image to base64
        image_data = base64.b64encode(uploaded_file.read()).decode('utf-8')
        
        # Create vision request
        response = client.chat.completions.create(
            model="gpt-4-vision-preview",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Please extract ALL text from this image, including any recipes, ingredients, instructions, or other text. Return the text exactly as it appears."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_data}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=2000
        )
        
        extracted_text = response.choices[0].message.content
        st.success(f"AI Vision extracted {len(extracted_text)} characters")
        return extracted_text
        
    except Exception as e:
        print(f"Vision API error: {e}")
        return ""

def extract_text_from_image(uploaded_file):
    try:
        # Reset file pointer
        uploaded_file.seek(0)
        
        # Try to open image
        try:
            image = Image.open(uploaded_file)
        except Exception as e:
            st.warning(f"Could not open image with PIL: {str(e)}")
            # Fall back to vision API immediately for unsupported formats
            st.info("Using AI vision for text extraction...")
            return extract_text_with_vision(uploaded_file)
        
        # Convert to RGB if necessary (handles RGBA, LA, etc.)
        if image.mode not in ('RGB', 'L'):
            try:
                image = image.convert('RGB')
            except Exception as e:
                st.warning(f"Could not convert image mode: {str(e)}")
                return extract_text_with_vision(uploaded_file)
        
        # Apply image preprocessing for better OCR
        # Resize if image is too small
        width, height = image.size
        if width < 1000:
            scale = 1000 / width
            new_width = int(width * scale)
            new_height = int(height * scale)
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Try OCR with different configurations
        try:
            # First try with default settings
            text = pytesseract.image_to_string(image)
            if text.strip():
                st.success(f"OCR extracted {len(text)} characters")
                return text
        except Exception as e:
            st.warning(f"Default OCR failed: {str(e)}")
            pass
        
        # Try with different PSM modes for better detection
        for psm in [6, 3, 11]:  # Different page segmentation modes
            try:
                custom_config = f'--oem 3 --psm {psm}'
                text = pytesseract.image_to_string(image, config=custom_config)
                if text.strip():
                    print(f"OCR succeeded with PSM mode {psm}")
                    return text
            except:
                continue
        
        st.warning("OCR failed to extract text. Trying AI vision method...")
        # Try vision API as last resort
        vision_text = extract_text_with_vision(uploaded_file)
        if vision_text:
            return vision_text
        return ""
        
    except Exception as e:
        st.error(f"Image processing error: {str(e)}")
        print(f"Image OCR error: {e}")
        return ""

def extract_text_from_csv(uploaded_file):
    try:
        text = uploaded_file.read().decode("utf-8", errors="ignore")
        uploaded_file.seek(0)
        reader = csv.reader(io.StringIO(text))
        return "\n".join([", ".join(row) for row in reader])
    except Exception as e:
        print(f"CSV parse error: {e}")
        return ""

def extract_text_from_docx(uploaded_file):
    try:
        document = Document(uploaded_file)
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        print(f"DOCX parse error: {e}")
        return ""

# --------------------------------------------
# 🖼️ Image Extraction Helpers
# --------------------------------------------

import base64
from urllib.parse import urljoin


def extract_image_from_pdf(uploaded_file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name
        doc = fitz.open(tmp_path)
        for page in doc:
            images = page.get_images(full=True)
            if images:
                xref = images[0][0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image.get("image")
                ext = base_image.get("ext", "png")
                b64 = base64.b64encode(img_bytes).decode("utf-8")
                doc.close()
                os.remove(tmp_path)
                return f"data:image/{ext};base64,{b64}"
        doc.close()
        os.remove(tmp_path)
    except Exception as e:
        print(f"PDF image extraction error: {e}")
    return None


def extract_image_from_docx_file(uploaded_file):
    try:
        document = Document(uploaded_file)
        for rel in document.part._rels.values():
            target = rel.target_part
            if "image" in target.content_type:
                b64 = base64.b64encode(target.blob).decode("utf-8")
                ext = target.content_type.split("/")[-1]
                return f"data:image/{ext};base64,{b64}"
    except Exception as e:
        print(f"DOCX image extraction error: {e}")
    return None


def extract_image_from_file(uploaded_file):
    """Upload the first discovered image to Firebase Storage and return its URL."""
    from firebase_init import get_bucket
    import uuid

    try:
        uploaded_file.seek(0)
        bucket = get_bucket()

        img_bytes = None
        ext = None
        if uploaded_file.type.startswith("image"):
            img_bytes = uploaded_file.read()
            ext = uploaded_file.type.split("/")[-1]
        elif uploaded_file.type == "application/pdf":
            data_url = extract_image_from_pdf(uploaded_file)
            if data_url:
                header, b64 = data_url.split(",", 1)
                img_bytes = base64.b64decode(b64)
                ext = header.split("/")[-1].split(";")[0]
        elif uploaded_file.type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            data_url = extract_image_from_docx_file(uploaded_file)
            if data_url:
                header, b64 = data_url.split(",", 1)
                img_bytes = base64.b64decode(b64)
                ext = header.split("/")[-1].split(";")[0]

        if img_bytes and ext:
            blob = bucket.blob(f"parsed_images/{uuid.uuid4()}.{ext}")
            blob.upload_from_string(img_bytes, content_type=f"image/{ext}")
            blob.make_public()
            return blob.public_url
    except Exception as e:
        print(f"Image extraction error: {e}")
    return None


def extract_image_from_soup(soup, base_url):
    props = ["og:image", "og:image:url", "twitter:image"]
    for prop in props:
        tag = soup.find("meta", property=prop)
        if tag and tag.get("content"):
            return urljoin(base_url, tag["content"])
    img_tag = soup.find("img")
    if img_tag and img_tag.get("src"):
        return urljoin(base_url, img_tag["src"])
    return None

# --------------------------------------------
# 🤖 AI Prompt Routing (Patched)
# --------------------------------------------

def query_ai_parser(raw_text, target_type):
    if not client:
        st.error("❌ OpenAI client not initialized. Please check your API key in .streamlit/secrets.toml")
        st.info("💡 Add your OpenAI API key to .streamlit/secrets.toml:\n[openai]\napi_key = \"sk-...\"")
        return {}
    
    system_prompt = (
        "You are an expert data parser. Extract only structured data from unstructured text.\n"
        "Return only a JSON object using proper capitalization.\n"
        "When parsing recipes, include a concise list of relevant tags such as cuisine, meal type, diets or allergens.\n"
        "- recipes → include name, ingredients, instructions, serves (number of servings), tags\n"
        "- menus → day, meal, items\n"
        "- tags → list of relevant tags\n"
        "- ingredients → list of items with quantity + unit if available\n"
        "- allergens → list of known allergens mentioned"
    )

    user_prompt = f"""
Extract structured {target_type} data from the following text. 
Use common section headers like:
- Ingredients:
- Instructions:
- Steps:
- Servings:

For recipes, ensure the JSON includes:
- "name": recipe title
- "ingredients": list of ingredients  
- "instructions": cooking steps
- "serves": number of servings (as a number, not string)
- "tags": relevant tags
- "allergens": any allergens mentioned

Only return a JSON object.

```
{raw_text[:6000]}
```
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )

        raw_output = response.choices[0].message.content
        # print("🔍 AI raw output:", raw_output)  # Debug only

        try:
            return json.loads(raw_output)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", raw_output, re.DOTALL)
            if match:
                return json.loads(match.group(0))
            raise

    except json.JSONDecodeError:
        st.error("❌ Failed to parse AI response as valid JSON.")
        return {}

    except Exception as e:
        error_msg = str(e)
        if "invalid_api_key" in error_msg or "401" in error_msg:
            # Extract the key that's being used from the error message
            import re
            key_match = re.search(r'provided: (\S+)\.', error_msg)
            if key_match:
                bad_key = key_match.group(1)
                st.error(f"❌ Invalid OpenAI API key being used: {bad_key[:20]}...")
                st.info(f"Expected key from secrets starts with: {st.secrets['openai']['api_key'][:20]}...")
            else:
                st.error("❌ Invalid OpenAI API key. Please check your configuration.")
            st.info("💡 Check .streamlit/secrets.toml and ensure no OPENAI_API_KEY environment variable is set.")
        else:
            st.error(f"OpenAI error: {e}")
        return {}

# --------------------------------------------
# 🌐 Parse Recipe From URL (Patched)
# --------------------------------------------

def parse_recipe_from_url(url: str) -> dict:
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        text = soup.get_text(separator="\n")
        image_url = extract_image_from_soup(soup, url)
    except requests.exceptions.RequestException as e:
        st.error(f"Failed to fetch page: {e}")
        if "403" in str(e) or "forbidden" in str(e).lower():
            st.info("💡 Tip: Some websites block automated requests. Try saving the recipe as a PDF and uploading it instead.")
        return {}
    except Exception as e:
        st.error(f"Error parsing URL: {e}")
        return {}

    cleaned_text = clean_raw_text(text)
    
    # If OpenAI is not available, provide manual entry option
    if not client:
        st.warning("⚠️ AI parsing is not available. Please use manual entry or upload a file.")
        return {}
    
    parsed = query_ai_parser(cleaned_text, "recipes")

    # API may return the recipe nested under a "recipes" key
    recipe = parsed
    if isinstance(parsed, dict) and "recipes" in parsed:
        recipe = parsed["recipes"]
        if isinstance(recipe, list):
            recipe = recipe[0] if recipe else {}

    # Normalize key casing for downstream logic
    recipe = normalize_keys(recipe)
    normalize_recipe_quantities(recipe)
    if image_url and not recipe.get("image_url"):
        recipe["image_url"] = image_url

    if not is_meaningful_recipe(recipe):
        st.warning(
            "⚠️ Recipe content extracted from URL appears to be incomplete or unusable."
        )
        return {}

    return recipe

# --------------------------------------------
# 📄 Parse Recipe From File
# --------------------------------------------

def parse_recipe_from_file(uploaded_file) -> dict:
    """Extract text from an uploaded file and parse the first recipe."""
    raw_text = extract_text(uploaded_file)
    if not raw_text or not raw_text.strip():
        return {}

    image_url = extract_image_from_file(uploaded_file)
    cleaned_text = clean_raw_text(raw_text)
    parsed = query_ai_parser(cleaned_text, "recipes")

    recipe = parsed
    if isinstance(parsed, dict) and "recipes" in parsed:
        recipe = parsed["recipes"]
        if isinstance(recipe, list):
            recipe = recipe[0] if recipe else {}

    recipe = normalize_keys(recipe)
    normalize_recipe_quantities(recipe)
    if image_url and not recipe.get("image_url"):
        recipe["image_url"] = image_url

    return recipe

# --------------------------------------------
# 💾 Modular Save Buttons
# --------------------------------------------

def render_extraction_buttons(file_id, parsed_data, user_id=None):
    from utils import get_active_event_id

    if not parsed_data:
        st.info("No parsed data available.")
        return

    st.markdown("## 💾 Save Extracted Content")

    if "recipes" in parsed_data and parsed_data["recipes"]:
        recipes = parsed_data["recipes"]
        recipes = recipes if isinstance(recipes, list) else [recipes]
        if st.button("📥 Save Recipes"):
            try:
                for r in recipes:
                    save_recipe_to_firestore(r, user_id=user_id, file_id=file_id)
                st.success("✅ Recipes saved to database")
            except Exception as e:
                st.error(f"Failed to save recipes: {e}")

    if "menus" in parsed_data and parsed_data["menus"]:
        menus = parsed_data["menus"]
        menus = menus if isinstance(menus, list) else [menus]
        event_id = get_active_event_id()
        if st.button("📥 Save Menus"):
            try:
                for m in menus:
                    save_menu_to_firestore(m, event_id=event_id, user_id=user_id, file_id=file_id)
                st.success("✅ Menus saved to database")
            except Exception as e:
                st.error(f"Failed to save menus: {e}")

    if "events" in parsed_data and parsed_data["events"]:
        events = parsed_data["events"]
        events = events if isinstance(events, list) else [events]
        if st.button("📥 Save Events"):
            try:
                for e in events:
                    save_event_to_firestore(e, user_id=user_id, file_id=file_id)
                st.success("✅ Events saved to database")
            except Exception as e:
                st.error(f"Failed to save events: {e}")

    if "ingredients" in parsed_data and parsed_data["ingredients"]:
        ingredients = parsed_data["ingredients"]
        if st.button("📥 Save Ingredients"):
            try:
                for ing in ingredients:
                    save_ingredient_to_firestore(ing, user_id=user_id, file_id=file_id)
                st.success("✅ Ingredients saved to database")
            except Exception as e:
                st.error(f"Failed to save ingredients: {e}")

    st.markdown("---")
    st.subheader("🧪 Parsed Data")
    st.json(parsed_data)
